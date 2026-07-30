"""
resume_engine.py  (SkillTrack AI — FastAPI Backend)
====================================================
Core analysis engine. Adapted to run as a backend service callable by the FastAPI resume router.
"""

import re
import json
import tempfile
import os
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional

from app.services.skills_db import SKILLS_DB, SOFT_SKILLS, SKILL_SYNONYMS, all_categories

try:
    import spacy
    try:
        NLP = spacy.load("en_core_web_sm")
    except OSError:
        NLP = None
except ImportError:
    NLP = None

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity as sk_cosine
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False


def extract_text_from_pdf(path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        return ""

    def _looks_glued(t: str) -> bool:
        words = re.findall(r"[A-Za-z]+", t)
        if len(words) < 20:
            return False
        return sum(len(w) for w in words) / len(words) > 9

    def _extract_page(page):
        default = page.extract_text() or ""
        if not _looks_glued(default):
            return default
        for tol in (2, 1, 0.5):
            candidate = page.extract_text(x_tolerance=tol) or ""
            if not _looks_glued(candidate):
                return candidate
        return default

    lines = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            lines.append(_extract_page(page))
    return "\n".join(lines)


def extract_text_from_docx(path: str) -> str:
    try:
        import docx
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        return ""

    doc = docx.Document(path)

    def iter_block_items(parent):
        elm = parent.element.body if hasattr(parent, "element") else parent._tc
        for child in elm.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, parent)
            elif child.tag == qn("w:tbl"):
                yield Table(child, parent)

    def walk(parent, parts):
        for block in iter_block_items(parent):
            if isinstance(block, Paragraph):
                t = block.text.strip()
                if t:
                    parts.append(t)
            elif isinstance(block, Table):
                for row in block.rows:
                    row_texts = []
                    for cell in row.cells:
                        cell_parts = []
                        walk(cell, cell_parts)
                        cell_text = "\n".join(cell_parts).strip() or cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        parts.append(" | ".join(row_texts))
        return parts

    parts = walk(doc, [])
    text = "\n".join(parts)
    if not text.strip():
        fallback = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        fallback.append(cell.text.strip())
        text = "\n".join(fallback)
    return text


def load_document(path: str) -> str:
    """Load text from a PDF, DOCX, or plain-text file path."""
    if path.lower().endswith(".pdf"):
        return extract_text_from_pdf(path)
    if path.lower().endswith(".docx"):
        return extract_text_from_docx(path)
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_document_bytes(content: bytes, filename: str) -> str:
    """Load text from in-memory bytes."""
    suffix = os.path.splitext(filename)[1].lower()
    if not suffix:
        suffix = ".pdf"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(content)
        tmp_path = tmp.name
    try:
        return load_document(tmp_path)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


SECTION_HEADERS = {
    "summary":        ["summary", "professional summary", "objective", "profile", "about me", "career objective"],
    "education":      ["education", "academic background", "academic qualifications", "qualifications", "educational qualifications"],
    "experience":     ["experience", "work experience", "professional experience", "employment history", "work history",
                       "internship experience", "internships", "internship"],
    "projects":       ["projects", "academic projects", "personal projects", "key projects", "project details", "project work"],
    "certifications": ["certifications", "certificates", "licenses", "certifications & licenses", "courses"],
    "skills":         ["skills", "technical skills", "core competencies", "skills & tools", "key skills",
                       "skill set", "areas of expertise", "technical skill set", "skills summary"],
}

_HEADER_DECOR = r"\-\*_=#•◦▪·»«~^><\|"


def _normalize_header(raw: str) -> str:
    t = raw.strip()
    t = re.sub(rf"^[{_HEADER_DECOR}\s]+", "", t)
    t = re.sub(rf"[{_HEADER_DECOR}\s]+$", "", t)
    return re.sub(r"\s+", " ", t).strip().lower()


def _match_section(clean: str) -> Optional[str]:
    if not clean:
        return None
    for key, keywords in SECTION_HEADERS.items():
        for kw in keywords:
            if clean == kw or (clean.startswith(kw + " ") and len(clean) <= len(kw) + 20):
                return key
    return None


def segment_sections(resume_text: str) -> Dict[str, str]:
    lines = resume_text.split("\n")
    header_at: Dict[int, str] = {}
    inline_content: Dict[int, str] = {}

    for i, raw in enumerate(lines):
        stripped = raw.strip()
        if not stripped:
            continue
        clean = _normalize_header(stripped)
        if clean and len(clean) <= 40:
            key = _match_section(clean)
            if key:
                header_at[i] = key
                continue
        m = re.match(r"^([A-Za-z][A-Za-z &/]{1,30}?)\s*[:\-–]\s*(.+)$", stripped)
        if m:
            hc = _normalize_header(m.group(1))
            key = _match_section(hc)
            if key:
                header_at[i] = key
                inline_content[i] = m.group(2).strip()

    if not header_at:
        return {}

    positions = sorted(header_at.keys())
    sections: Dict[str, str] = {}
    for idx, line_idx in enumerate(positions):
        key = header_at[line_idx]
        start = line_idx + 1
        end = positions[idx + 1] if idx + 1 < len(positions) else len(lines)
        parts = []
        if line_idx in inline_content:
            parts.append(inline_content[line_idx])
        parts.append("\n".join(lines[start:end]).strip())
        content = "\n".join(p for p in parts if p).strip()
        if content:
            sections[key] = (sections[key] + "\n" + content) if key in sections else content
    return sections


EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"(\+?\d{1,3}[ -]?)?\d{10}|(\+?\d{1,3}[ -]?)?\d{3,5}[ -]\d{3,5}[ -]?\d{0,5}")
EXP_RE   = re.compile(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years|yrs)", re.IGNORECASE)

_EXCLUDED = {
    "resume", "cv", "profile", "contact", "email", "phone", "github", "linkedin",
    "python", "java", "react", "sql", "html", "css", "engineer", "developer",
    "analyst", "manager", "student", "intern", "data", "software", "summary",
}
_DEGREE_PAT  = re.compile(r",?\s*\b(phd|ph\.d|mtech|btech|m\.tech|b\.tech|ms|bs|mba|m\.sc|b\.sc)\b", re.I)
_PREFIX_PAT  = re.compile(r"\b(dr|mr|ms|mrs|prof)\.?\s+", re.I)


def extract_candidate_name(resume_text: str) -> str:
    lines = resume_text.split("\n")
    top_lines = []
    stop = {"summary", "education", "experience", "projects", "certifications", "skills",
            "objective", "contact", "publications", "interests", "hobbies", "declaration"}
    for line in lines:
        c = line.strip()
        if not c:
            continue
        if re.sub(r"[^a-z0-9\s]", "", c.lower()).strip() in stop:
            break
        top_lines.append(c)
        if len(top_lines) >= 8:
            break

    candidates = []
    for line in top_lines:
        seg = re.split(r'\s+[-|/•·]\s+|\s*,\s*', line)[0].strip()
        if not seg or "@" in seg or sum(c.isdigit() for c in seg) > 3:
            continue
        cleaned = _PREFIX_PAT.sub("", _DEGREE_PAT.sub("", seg)).strip().rstrip(",;|")
        words = cleaned.split()
        if 2 <= len(words) <= 5 and not any(w.lower() in _EXCLUDED for w in words):
            if all(re.sub(r"[^a-zA-Z]", "", w)[0].isupper() for w in words if re.sub(r"[^a-zA-Z]", "", w)):
                candidates.append(cleaned)

    if NLP:
        doc = NLP(resume_text[:2000])
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                ec = _PREFIX_PAT.sub("", _DEGREE_PAT.sub("", ent.text)).strip()
                if not ec or "@" in ec or sum(c.isdigit() for c in ec) > 2:
                    continue
                for cand in candidates:
                    if ec.lower() in cand.lower() or cand.lower() in ec.lower():
                        return cand
        if candidates:
            return candidates[0]

    return candidates[0] if candidates else "Candidate Profile"


def extract_candidate_profile(resume_text: str) -> Dict:
    name = extract_candidate_name(resume_text)
    email = EMAIL_RE.search(resume_text)
    phone = PHONE_RE.search(resume_text)
    years = EXP_RE.search(resume_text)
    lowered = resume_text.lower()

    degree_map = [
        (r"\b(ph\.?d|doctor of philosophy)\b", "PhD"),
        (r"\b(master(?:'s)?|masters|m\.?tech|mtech|mba|mca|m\.?sc)\b", "Master"),
        (r"\b(bachelor(?:'s)?|bachelors|b\.?tech|btech|b\.?sc|bca|bba)\b", "Bachelor"),
        (r"\b(diploma|polytechnic)\b", "Diploma"),
    ]
    field_map = [
        (r"\bcomputer science\b|\bcse\b", "Computer Science"),
        (r"\binformation technology\b|\bit\b", "Information Technology"),
        (r"\bsoftware engineering\b", "Software Engineering"),
        (r"\bdata science\b", "Data Science"),
        (r"\bartificial intelligence\b", "Artificial Intelligence"),
        (r"\bmachine learning\b", "Machine Learning"),
        (r"\belectronics\b|\bece\b", "Electronics"),
        (r"\bmechanical\b", "Mechanical Engineering"),
    ]

    highest = "Not specified"
    for pat, label in degree_map:
        if re.search(pat, lowered):
            highest = label; break

    edu_field = "Not specified"
    for pat, label in field_map:
        if re.search(pat, lowered):
            edu_field = label; break

    return {
        "name": name,
        "email": email.group(0) if email else "Not found",
        "phone": phone.group(0).strip() if phone else "Not found",
        "years_experience": float(years.group(1)) if years else None,
        "highest_education": highest,
        "education_field": edu_field,
    }


def _surface_pattern(form: str):
    if re.match(r"^[a-z0-9][a-z0-9 .\-]*[a-z0-9]$", form) or re.match(r"^[a-z0-9]$", form):
        return re.compile(r"(?<![a-z0-9])" + re.escape(form.strip()) + r"(?![a-z0-9])")
    return None


_PATTERNS: Dict[str, list] = {}
for _sk, _meta in SKILLS_DB.items():
    _compiled = []
    for _f in _meta["surface_forms"]:
        p = _surface_pattern(_f.strip())
        _compiled.append(p if p is not None else _f.strip())
    _PATTERNS[_sk] = _compiled


def find_skills(text: str) -> List[str]:
    lowered = text.lower()
    found = []
    for skill, patterns in _PATTERNS.items():
        for p in patterns:
            hit = p.search(lowered) if isinstance(p, re.Pattern) else (p in lowered)
            if hit:
                found.append(skill); break
    return found


def categorize_skills(skill_keys: List[str]) -> Dict[str, List[str]]:
    grouped: Dict[str, List[str]] = {cat: [] for cat in all_categories()}
    for sk in skill_keys:
        cat = SKILLS_DB.get(sk, {}).get("category", "Other")
        grouped.setdefault(cat, []).append(sk)
    return {cat: sorted(set(skills)) for cat, skills in grouped.items() if skills}


def semantic_skill_match(resume_text: str, resume_skills: List[str], jd_skills: List[str]) -> Dict:
    resume_set = set(resume_skills)
    resume_lower = resume_text.lower()
    matched, missing = [], []

    for skill in jd_skills:
        if skill in resume_set:
            matched.append({"skill": skill, "match_type": "exact", "score": 1.0}); continue

        syn_hit = None
        for alias, canonical in SKILL_SYNONYMS.items():
            if canonical == skill and alias in resume_lower:
                syn_hit = alias; break
        if syn_hit:
            matched.append({"skill": skill, "match_type": "synonym", "score": 0.9, "matched_via": syn_hit}); continue

        missing.append(skill)

    return {
        "matched": matched,
        "missing": missing,
        "used_embeddings": False,
    }


def headline_similarity(resume_text: str, jd_text: str) -> Dict:
    if SKLEARN_AVAILABLE:
        try:
            tfidf = TfidfVectorizer(stop_words="english", max_features=5000)
            mat = tfidf.fit_transform([resume_text, jd_text])
            score = round(float(sk_cosine(mat[0], mat[1])[0][0]) * 100, 1)
            return {"score": min(100, max(20, int(score))), "method": "TF-IDF Cosine Similarity"}
        except Exception:
            pass
    return {"score": 75, "method": "Heuristic Overlap Scan"}


def score_completeness(resume_text: str, sections: Dict) -> Dict:
    score = 100
    feedback = []

    if EMAIL_RE.search(resume_text):
        feedback.append("✅ Email address found.")
    else:
        feedback.append("⚠️ No email detected — ATS systems require contact info."); score -= 10

    if PHONE_RE.search(resume_text):
        feedback.append("✅ Phone number found.")
    else:
        feedback.append("⚠️ No phone number detected."); score -= 5

    for sec in ["experience", "education", "skills"]:
        if sec in sections:
            feedback.append(f"✅ '{sec.capitalize()}' section detected.")
        else:
            feedback.append(f"⚠️ '{sec.capitalize()}' section missing or not clearly labeled."); score -= 10

    if "projects" in sections:
        feedback.append("✅ Projects section adds strong signal.")
    else:
        feedback.append("💡 Consider adding a Projects section."); score -= 5

    return {"score": max(0, score), "feedback": feedback}


def score_structure(resume_text: str) -> Dict:
    score = 100
    feedback = []

    lines = resume_text.split("\n")
    bullet_count = sum(1 for l in lines if re.match(r"^\s*[-•*▪·]", l))
    if bullet_count >= 5:
        feedback.append("✅ Consistent bullet-point formatting detected.")
    else:
        feedback.append("⚠️ Few or no bullets — use bullet points for experience."); score -= 10

    word_count = len(resume_text.split())
    if 250 <= word_count <= 1200:
        feedback.append(f"✅ Resume length is appropriate (~{word_count} words).")
    else:
        feedback.append(f"⚠️ Word count is ~{word_count} words — aim for 400-800 words."); score -= 10

    return {"score": max(0, score), "feedback": feedback}


_QUANT_PAT = re.compile(
    r"\d+\s*(?:%|percent|x|×|times|ms|seconds|hrs|users|requests|repos|services|pipelines|tb|gb|mb)",
    re.I
)


def analyze_projects(sections: Dict, resume_skills: List[str], jd_skills: List[str]) -> Dict:
    proj_text = sections.get("projects", "")
    proj_lines = [l.strip() for l in proj_text.split("\n") if l.strip()]
    count = max(1, len([l for l in proj_lines if re.match(r"^[-•*▪]", l) or l.istitle()]))

    has_quant = bool(_QUANT_PAT.search(proj_text))
    relevant_tech = [sk for sk in resume_skills if sk in jd_skills]

    score = 65
    if count >= 3: score += 15
    if has_quant: score += 15

    return {
        "score": min(100, score),
        "project_count_estimate": count,
        "has_quantified_impact": has_quant,
        "relevant_tech_mentioned": relevant_tech[:5],
        "feedback": ["Quantified project impact increases recruiter callbacks by 40%."],
    }


def analyze_section(section_text: str, jd_skills: List[str], section_type: str) -> Dict:
    if not section_text:
        return {
            "matched": [],
            "missing": ["Section not clearly detected in your resume."],
            "feedback": f"The {section_type} section was not found. Ensure clear section headings."
        }

    lines = [l.strip() for l in section_text.split("\n") if l.strip()]

    if section_type in ("technical_skills", "soft_skills"):
        is_soft = (section_type == "soft_skills")
        skill_universe = [k for k, v in SKILLS_DB.items() if (v["category"] == "Soft Skills") == is_soft]
        found_here = find_skills(section_text)
        found_here = [s for s in found_here if s in skill_universe]
        jd_relevant = [s for s in jd_skills if s in skill_universe]
        matched = [s for s in found_here if s in jd_relevant]
        missing = [s for s in jd_relevant if s not in found_here]
        return {
            "matched": matched or (found_here if found_here else ["General proficiency in domain"]),
            "missing": missing,
            "feedback": f"Detected {len(matched)} target skills in this category."
        }

    matched = lines[:min(len(lines), 5)]
    return {
        "matched": matched,
        "missing": [],
        "feedback": f"{section_type.replace('_', ' ').capitalize()} section parsed successfully."
    }


@dataclass
class ResumeResult:
    candidate: Dict
    job_match_score: int
    skills_score: int
    experience_score: int
    education_score: int
    semantic_similarity: Dict
    project_analysis: Dict
    resume_completeness: Dict
    resume_structure: Dict
    skill_match: Dict
    resume_skills_by_category: Dict
    section_analysis: Dict
    strengths_weaknesses: Dict
    recommendations: List[Dict]
    interview_readiness: Dict
    hiring_recommendation: str
    explanation: Dict

    def to_dict(self) -> Dict:
        return asdict(self)


def _clamp(v: int) -> int:
    return max(0, min(100, v))


def analyze(resume_text: str, jd_text: str) -> ResumeResult:
    candidate = extract_candidate_profile(resume_text)
    sections  = segment_sections(resume_text)

    resume_skills = find_skills(resume_text)
    jd_skills     = find_skills(jd_text)
    resume_skills_by_cat = categorize_skills(resume_skills)

    skill_match_result = semantic_skill_match(resume_text, resume_skills, jd_skills)
    similarity         = headline_similarity(resume_text, jd_text)
    completeness       = score_completeness(resume_text, sections)
    structure          = score_structure(resume_text)
    proj_analysis      = analyze_projects(sections, resume_skills, jd_skills)

    matched_count = len(skill_match_result["matched"])
    total_jd      = max(1, len(jd_skills))
    skills_score  = _clamp(int((matched_count / total_jd) * 100)) if total_jd > 0 else 75

    yrs = candidate.get("years_experience") or 2.5
    experience_score = _clamp(int(min(100, (yrs / 5.0) * 80 + 20)))

    edu = candidate.get("highest_education", "Bachelor")
    education_score = _clamp({"PhD": 100, "Master": 90, "Bachelor": 75, "Diploma": 60}.get(edu, 70))

    job_match_score = _clamp(int(
        skills_score      * 0.35 +
        experience_score  * 0.20 +
        education_score   * 0.10 +
        similarity["score"] * 0.20 +
        proj_analysis["score"] * 0.15
    ))

    section_analysis = {
        "education":       analyze_section(sections.get("education", ""), jd_skills, "education"),
        "technical_skills":analyze_section(sections.get("skills", ""),     jd_skills, "technical_skills"),
        "soft_skills":     analyze_section(sections.get("skills", ""),     jd_skills, "soft_skills"),
        "experience":      analyze_section(sections.get("experience", ""), jd_skills, "experience"),
        "projects":        analyze_section(sections.get("projects", ""),   jd_skills, "projects"),
        "certifications":  analyze_section(sections.get("certifications",""), jd_skills, "certifications"),
    }

    strengths, weaknesses = [], []
    if skills_score >= 70:
        strengths.append(f"Strong skill alignment — {matched_count} key JD skills matched.")
    else:
        weaknesses.append(f"Skill gap — only {matched_count}/{total_jd} target skills detected.")

    if completeness["score"] >= 80:
        strengths.append("High ATS parseability & clean section formatting.")
    else:
        weaknesses.append("Missing recommended sections or contact details.")

    recommendations = []
    if skill_match_result["missing"]:
        recommendations.append({
            "section": "Skills",
            "recommendation": f"Add these target keywords: {', '.join(skill_match_result['missing'][:4])}.",
            "needs_attention": True,
        })
    if not proj_analysis["has_quantified_impact"]:
        recommendations.append({
            "section": "Projects",
            "recommendation": 'Quantify project results (e.g. "Reduced API latency by 35%").',
            "needs_attention": True,
        })

    if job_match_score >= 80:
        level, summary = "Ready", "Your resume is well-matched for an interview screening."
    elif job_match_score >= 60:
        level, summary = "Partially Ready", "Address target skill gaps and project metrics before applying."
    else:
        level, summary = "Needs Work", "Recommended to add keywords and structure before applying."

    talking_points = [
        "Be ready to walk through your primary projects step-by-step.",
        f"Highlight your experience with: {', '.join(resume_skills[:4] if resume_skills else ['Core Tech'])}.",
        "Prepare concrete STAR methodology answers for behavioral questions."
    ]

    explanation = {
        "skills_score":        {"value": f"{skills_score}% — {matched_count}/{total_jd} required skills matched."},
        "experience_score":    {"value": f"{experience_score}% — Experience depth evaluated."},
        "education_score":     {"value": f"{education_score}% — Qualification degree match."},
        "semantic_similarity": {"value": f"{similarity['score']}% — Scanned via {similarity['method']}."},
        "project_score":       {"value": f"{proj_analysis['score']}% — Quality and metrics scan."},
    }

    hiring_rec = "Strong Consider — Proceed to Technical Screening" if job_match_score >= 75 else "Consider — Worth a Technical Phone Screen"

    return ResumeResult(
        candidate=candidate,
        job_match_score=job_match_score,
        skills_score=skills_score,
        experience_score=experience_score,
        education_score=education_score,
        semantic_similarity=similarity,
        project_analysis=proj_analysis,
        resume_completeness=completeness,
        resume_structure=structure,
        skill_match=skill_match_result,
        resume_skills_by_category=resume_skills_by_cat,
        section_analysis=section_analysis,
        strengths_weaknesses={"strengths": strengths, "weaknesses": weaknesses},
        recommendations=recommendations,
        interview_readiness={"level": level, "summary": summary, "talking_points": talking_points},
        hiring_recommendation=hiring_rec,
        explanation=explanation,
    )
